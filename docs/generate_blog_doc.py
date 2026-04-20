#!/usr/bin/env python3
"""
Generate Word Document from Technical Blog
With PlantUML diagrams converted to images
"""

import os
import re
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
import numpy as np

# Set font for matplotlib to support Chinese
plt.rcParams['font.sans-serif'] = ['Microsoft YaHei', 'SimHei', 'DejaVu Sans']
plt.rcParams['axes.unicode_minus'] = False

def add_heading(doc, text, level):
    """Add a heading with specified level"""
    heading = doc.add_heading(text, level=level)
    return heading

def add_paragraph(doc, text, bold=False, italic=False):
    """Add a paragraph with optional formatting"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    return para

def add_code_block(doc, code, language=""):
    """Add a code block with syntax highlighting effect"""
    para = doc.add_paragraph()
    para.style = 'No Spacing'
    run = para.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x2E, 0x86, 0xAB)
    # Set Chinese font
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return para

def add_image(doc, image_path, width=None, height=None):
    """Add an image to the document"""
    if os.path.exists(image_path):
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        if width:
            run.add_picture(image_path, width=width)
        elif height:
            run.add_picture(image_path, height=height)
        else:
            run.add_picture(image_path)
        return para
    return None

def create_sdd_evolution_diagram():
    """Create SDD evolution diagram"""
    fig, ax = plt.subplots(1, 1, figsize=(10, 5))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 6)
    ax.axis('off')
    ax.set_title('SDD三次认知迭代', fontsize=14, fontweight='bold', pad=20)

    # Draw boxes for each evolution
    evolutions = [
        (1.5, 3, '2019年\nDDD\n领域驱动设计', '#E8D5B7'),
        (5, 3, '2022年\nAPI-First\n接口契约驱动', '#B8D4E3'),
        (8.5, 3, '2024年\nLLM-Based SDD\n自然语言驱动', '#C5E0B4'),
    ]

    for x, y, text, color in evolutions:
        box = FancyBboxPatch((x-1.2, y-0.8), 2.4, 1.8,
                            boxstyle="round,pad=0.05,rounding_size=0.2",
                            facecolor=color, edgecolor='#333333', linewidth=2)
        ax.add_patch(box)
        ax.text(x, y, text, ha='center', va='center', fontsize=10, fontweight='bold')

    # Draw arrows
    ax.annotate('', xy=(3.3, 3), xytext=(2.7, 3),
                arrowprops=dict(arrowstyle='->', color='#333333', lw=2))
    ax.annotate('', xy=(6.8, 3), xytext=(6.2, 3),
                arrowprops=dict(arrowstyle='->', color='#333333', lw=2))

    # Add labels
    ax.text(5, 1, '从"设计文档"到"Prompt契约"，复杂度从编码转移到描述', ha='center', fontsize=9, style='italic')

    plt.tight_layout()
    path = 'docs/images/sdd_evolution.png'
    os.makedirs('docs/images', exist_ok=True)
    plt.savefig(path, dpi=150, bbox_inches='tight', facecolor='white')
    plt.close()
    return path

def create_llm_provider_diagram():
    """Create LLM Provider abstraction diagram"""
    fig, ax = plt.subplots(1, 1, figsize=(10, 6))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 8)
    ax.axis('off')
    ax.set_title('多LLM Provider抽象设计', fontsize=14, fontweight='bold', pad=20)

    # Draw LLMClient abstract layer
    client_box = FancyBboxPatch((3.5, 6), 3, 1.2,
                                boxstyle="round,pad=0.05,rounding_size=0.15",
                                facecolor='#FFE4B5', edgecolor='#333333', linewidth=2)
    ax.add_patch(client_box)
    ax.text(5, 6.6, 'LLMClient\n(抽象接口)', ha='center', va='center', fontsize=11, fontweight='bold')

    # Draw provider boxes
    providers = [
        (1.5, 4, 'SQLCoder\nProvider', '#90EE90', '当前生产'),
        (5, 4, 'OpenAI\nProvider', '#87CEEB', '备选方案'),
        (8.5, 4, 'MiniMax\nProvider', '#DDA0DD', '成本敏感'),
    ]

    for x, y, text, color, note in providers:
        box = FancyBboxPatch((x-1.3, y-0.7), 2.6, 1.6,
                            boxstyle="round,pad=0.05,rounding_size=0.1",
                            facecolor=color, edgecolor='#333333', linewidth=2)
        ax.add_patch(box)
        ax.text(x, y+0.2, text, ha='center', va='center', fontsize=10, fontweight='bold')
        ax.text(x, y-0.4, note, ha='center', va='center', fontsize=8, style='italic')

    # Draw connecting lines
    for x in [2.8, 5, 7.2]:
        ax.annotate('', xy=(x, 5.2), xytext=(5, 6),
                   arrowprops=dict(arrowstyle='->', color='#666666', lw=1.5))

    # Draw core logic box
    core_box = FancyBboxPatch((3, 1), 4, 1.5,
                             boxstyle="round,pad=0.05,rounding_size=0.1",
                             facecolor='#F0F0F0', edgecolor='#333333', linewidth=2)
    ax.add_patch(core_box)
    ax.text(5, 1.75, '核心业务逻辑\n(不依赖具体LLM)', ha='center', va='center', fontsize=10)

    # Draw vertical line to providers
    ax.annotate('', xy=(5, 4.7), xytext=(5, 2.5),
               arrowprops=dict(arrowstyle='->', color='#666666', lw=1.5))

    plt.tight_layout()
    path = 'docs/images/llm_provider.png'
    os.makedirs('docs/images', exist_ok=True)
    plt.savefig(path, dpi=150, bbox_inches='tight', facecolor='white')
    plt.close()
    return path

def create_intent_detection_diagram():
    """Create intent detection cascade architecture diagram"""
    fig, ax = plt.subplots(1, 1, figsize=(10, 7))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 9)
    ax.axis('off')
    ax.set_title('意图识别级联设计 - LLM与规则引擎协作', fontsize=14, fontweight='bold', pad=20)

    # Draw input box
    input_box = FancyBboxPatch((3.5, 7.5), 3, 0.8,
                               boxstyle="round,pad=0.05,rounding_size=0.1",
                               facecolor='#FFE4E1', edgecolor='#333333', linewidth=2)
    ax.add_patch(input_box)
    ax.text(5, 7.9, '用户查询', ha='center', va='center', fontsize=10, fontweight='bold')

    # Draw rule engine layer
    rule_box = FancyBboxPatch((1, 5.5), 3.5, 1.8,
                              boxstyle="round,pad=0.05,rounding_size=0.1",
                              facecolor='#E6E6FA', edgecolor='#333333', linewidth=2)
    ax.add_patch(rule_box)
    ax.text(2.75, 6.6, '规则引擎匹配层', ha='center', va='center', fontsize=10, fontweight='bold')
    ax.text(2.75, 5.9, '• 高频模式（有多少）\n• 明确意图（查站点）\n• 正则/关键词', ha='center', va='center', fontsize=8)

    # Draw LLM detection
    llm_box = FancyBboxPatch((5.5, 5.5), 3.5, 1.8,
                             boxstyle="round,pad=0.05,rounding_size=0.1",
                             facecolor='#B0E0E6', edgecolor='#333333', linewidth=2)
    ax.add_patch(llm_box)
    ax.text(7.25, 6.6, 'LLM意图检测兜底', ha='center', va='center', fontsize=10, fontweight='bold')
    ax.text(7.25, 5.9, '• 处理模糊查询\n• 语义理解\n• 泛化能力', ha='center', va='center', fontsize=8)

    # Draw diamond for decision
    decision = plt.Polygon([(5, 4.3), (5.5, 3.8), (5, 3.3), (4.5, 3.8)],
                          facecolor='#FFD700', edgecolor='#333333', linewidth=2)
    ax.add_patch(decision)
    ax.text(5, 3.8, '匹配\n成功?', ha='center', va='center', fontsize=8, fontweight='bold')

    # Draw output boxes
    out1_box = FancyBboxPatch((0.5, 1.5), 3, 1.2,
                              boxstyle="round,pad=0.05,rounding_size=0.1",
                              facecolor='#90EE90', edgecolor='#333333', linewidth=2)
    ax.add_patch(out1_box)
    ax.text(2, 2.1, '直接路由\n可解释 ✓', ha='center', va='center', fontsize=10, fontweight='bold')

    out2_box = FancyBboxPatch((5.5, 1.5), 3.5, 1.2,
                              boxstyle="round,pad=0.05,rounding_size=0.1",
                              facecolor='#FFA500', edgecolor='#333333', linewidth=2)
    ax.add_patch(out2_box)
    ax.text(7.25, 2.1, 'LLM兜底处理\n模糊查询', ha='center', va='center', fontsize=10, fontweight='bold')

    # Draw arrows
    ax.annotate('', xy=(2.75, 7.3), xytext=(5, 7.5),
               arrowprops=dict(arrowstyle='->', color='#333333', lw=1.5))
    ax.annotate('', xy=(2.75, 7.3), xytext=(7.25, 7.5),
               arrowprops=dict(arrowstyle='->', color='#333333', lw=1.5))

    ax.annotate('', xy=(5, 4.3), xytext=(3, 5.5),
               arrowprops=dict(arrowstyle='->', color='#333333', lw=1.5))
    ax.annotate('', xy=(5, 4.3), xytext=(7.25, 5.5),
               arrowprops=dict(arrowstyle='->', color='#333333', lw=1.5))

    ax.annotate('', xy=(2, 2.7), xytext=(4.6, 3.4),
               arrowprops=dict(arrowstyle='->', color='#333333', lw=1.5))
    ax.annotate('', xy=(7.25, 2.7), xytext=(5.4, 3.4),
               arrowprops=dict(arrowstyle='->', color='#333333', lw=1.5))

    # Labels
    ax.text(1.5, 4.5, 'Yes', fontsize=9, style='italic')
    ax.text(8, 4.5, 'No', fontsize=9, style='italic')

    plt.tight_layout()
    path = 'docs/images/intent_detection.png'
    os.makedirs('docs/images', exist_ok=True)
    plt.savefig(path, dpi=150, bbox_inches='tight', facecolor='white')
    plt.close()
    return path

def create_rag_knowledge_diagram():
    """Create RAG knowledge enhancement diagram"""
    fig, ax = plt.subplots(1, 1, figsize=(10, 6))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 8)
    ax.axis('off')
    ax.set_title('RAG知识增强四步法', fontsize=14, fontweight='bold', pad=20)

    # Draw knowledge pyramid
    pyramid_points = [
        (5, 7, 3, 0.8, '#FFE4B5', '文档知识\nSchema/API'),
        (5, 5.8, 3.6, 0.8, '#B0E0E6', '领域知识\n业务规则/计算'),
        (5, 4.6, 4.2, 0.8, '#98FB98', '系统知识\n接口契约'),
        (5, 3.4, 4.8, 0.8, '#DDA0DD', '专业知识\n运营商知识'),
    ]

    for x, y, w, h, color, text in pyramid_points:
        box = FancyBboxPatch((x-w/2, y-h/2), w, h,
                            boxstyle="round,pad=0.03,rounding_size=0.05",
                            facecolor=color, edgecolor='#333333', linewidth=2)
        ax.add_patch(box)
        ax.text(x, y, text, ha='center', va='center', fontsize=9, fontweight='bold')

    # Add step numbers
    steps = ['Step 1', 'Step 2', 'Step 3', 'Step 4']
    for i, (x, y, w, h, _, _) in enumerate(pyramid_points):
        ax.text(x-w/2+0.3, y+h/2-0.15, steps[i], fontsize=8, fontweight='bold', color='#666')

    # Draw process flow
    ax.text(8, 7, '知识分层', fontsize=10, fontweight='bold', ha='center')
    ax.text(8, 5.8, '向量化策略', fontsize=10, fontweight='bold', ha='center')
    ax.text(8, 4.6, '检索增强时机', fontsize=10, fontweight='bold', ha='center')
    ax.text(8, 3.4, '效果验证', fontsize=10, fontweight='bold', ha='center')

    # Draw arrow down the right side
    for y in [6.4, 5.2, 4, 2.8]:
        ax.annotate('', xy=(8.5, y-0.2), xytext=(8.5, y+0.2),
                   arrowprops=dict(arrowstyle='->', color='#666', lw=1.5))

    plt.tight_layout()
    path = 'docs/images/rag_knowledge.png'
    os.makedirs('docs/images', exist_ok=True)
    plt.savefig(path, dpi=150, bbox_inches='tight', facecolor='white')
    plt.close()
    return path

def create_sql_safety_diagram():
    """Create SQL safety validation framework diagram"""
    fig, ax = plt.subplots(1, 1, figsize=(10, 7))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 9)
    ax.axis('off')
    ax.set_title('LLM生成内容安全校验框架', fontsize=14, fontweight='bold', pad=20)

    # Draw SQL input
    sql_input = FancyBboxPatch((0.5, 7), 4, 1,
                               boxstyle="round,pad=0.05,rounding_size=0.1",
                               facecolor='#FFE4E1', edgecolor='#FF6B6B', linewidth=2)
    ax.add_patch(sql_input)
    ax.text(2.5, 7.5, 'LLM生成的SQL', ha='center', va='center', fontsize=10, fontweight='bold')

    # Draw validation layers
    layers = [
        (2.5, 5.5, '语法安全层', '#E6E6FA', '• SELECT开头\n• LIMIT限制\n• 禁止危险关键词'),
        (2.5, 3.5, '业务安全层', '#B0E0E6', '• 表权限检查\n• 跨运营商隔离\n• 结果集限制'),
        (2.5, 1.5, '语义校验层', '#98FB98', '• 执行计划分析\n• 结果合理性检查\n• 异常模式检测'),
    ]

    for x, y, title, color, content in layers:
        box = FancyBboxPatch((x-2.2, y-1), 4.4, 2,
                            boxstyle="round,pad=0.05,rounding_size=0.1",
                            facecolor=color, edgecolor='#333333', linewidth=2)
        ax.add_patch(box)
        ax.text(x, y+0.8, title, ha='center', va='center', fontsize=11, fontweight='bold')
        ax.text(x, y, content, ha='center', va='center', fontsize=8)

    # Draw arrows between layers
    for y in [4.5, 2.5]:
        ax.annotate('', xy=(2.5, y), xytext=(2.5, y+1),
                   arrowprops=dict(arrowstyle='->', color='#333333', lw=2))

    # Draw output
    output_pass = FancyBboxPatch((6, 4), 3, 1.2,
                                boxstyle="round,pad=0.05,rounding_size=0.1",
                                facecolor='#90EE90', edgecolor='#333333', linewidth=2)
    ax.add_patch(output_pass)
    ax.text(7.5, 4.6, '执行SQL', ha='center', va='center', fontsize=10, fontweight='bold')

    output_fail = FancyBboxPatch((6, 1.5), 3, 1,
                                 boxstyle="round,pad=0.05,rounding_size=0.1",
                                 facecolor='#FF6B6B', edgecolor='#333333', linewidth=2)
    ax.add_patch(output_fail)
    ax.text(7.5, 2, '友好错误', ha='center', va='center', fontsize=10, fontweight='bold')

    # Draw decision arrow
    ax.annotate('', xy=(4.7, 4.6), xytext=(2.5, 5.5),
               arrowprops=dict(arrowstyle='->', color='#333333', lw=1.5))
    ax.annotate('', xy=(4.7, 2), xytext=(2.5, 1.5),
               arrowprops=dict(arrowstyle='->', color='#FF6B6B', lw=1.5))
    ax.annotate('', xy=(7.5, 4.6), xytext=(6, 5),
               arrowprops=dict(arrowstyle='->', color='#333333', lw=1.5))

    ax.text(5.5, 5.5, '通过', fontsize=9, color='#90EE90')
    ax.text(4.2, 3.5, '失败', fontsize=9, color='#FF6B6B')

    plt.tight_layout()
    path = 'docs/images/sql_safety.png'
    os.makedirs('docs/images', exist_ok=True)
    plt.savefig(path, dpi=150, bbox_inches='tight', facecolor='white')
    plt.close()
    return path

def create_multi_agent_diagram():
    """Create Multi-Agent collaboration diagram"""
    fig, ax = plt.subplots(1, 1, figsize=(10, 7))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 9)
    ax.axis('off')
    ax.set_title('Multi-Agent协作模式', fontsize=14, fontweight='bold', pad=20)

    # Draw question box
    question = FancyBboxPatch((3, 7.5), 4, 1,
                               boxstyle="round,pad=0.05,rounding_size=0.1",
                               facecolor='#FFE4E1', edgecolor='#333333', linewidth=2)
    ax.add_patch(question)
    ax.text(5, 8, '用户问题：北京联通NR频段\n4月份峰值速率是多少？', ha='center', va='center', fontsize=9, fontweight='bold')

    # Draw Planner Agent
    planner = FancyBboxPatch((3, 5.5), 4, 1.5,
                            boxstyle="round,pad=0.05,rounding_size=0.1",
                            facecolor='#FFD700', edgecolor='#333333', linewidth=2)
    ax.add_patch(planner)
    ax.text(5, 6.5, 'Planner Agent', ha='center', va='center', fontsize=11, fontweight='bold')
    ax.text(5, 5.9, '• 拆解问题\n• 并行分发任务', ha='center', va='center', fontsize=8)

    # Draw Specialist Agents
    specialists = [
        (1.5, 3, '指标专家\nAgent', '#87CEEB'),
        (5, 3, '知识库专家\nAgent', '#98FB98'),
        (8.5, 3, 'SQL专家\nAgent', '#DDA0DD'),
    ]

    for x, y, text, color in specialists:
        box = FancyBboxPatch((x-1.3, y-0.7), 2.6, 1.6,
                            boxstyle="round,pad=0.05,rounding_size=0.1",
                            facecolor=color, edgecolor='#333333', linewidth=2)
        ax.add_patch(box)
        ax.text(x, y, text, ha='center', va='center', fontsize=9, fontweight='bold')

    # Draw Verifier Agent
    verifier = FancyBboxPatch((3, 0.5), 4, 1.2,
                              boxstyle="round,pad=0.05,rounding_size=0.1",
                              facecolor='#FFA500', edgecolor='#333333', linewidth=2)
    ax.add_patch(verifier)
    ax.text(5, 1.1, 'Verifier Agent\n汇总验证', ha='center', va='center', fontsize=10, fontweight='bold')

    # Draw arrows
    ax.annotate('', xy=(5, 7), xytext=(5, 8),
               arrowprops=dict(arrowstyle='->', color='#333333', lw=2))

    ax.annotate('', xy=(1.5, 4.5), xytext=(4, 5.5),
               arrowprops=dict(arrowstyle='->', color='#333333', lw=1.5))
    ax.annotate('', xy=(5, 4.5), xytext=(5, 5.5),
               arrowprops=dict(arrowstyle='->', color='#333333', lw=1.5))
    ax.annotate('', xy=(8.5, 4.5), xytext=(6, 5.5),
               arrowprops=dict(arrowstyle='->', color='#333333', lw=1.5))

    ax.annotate('', xy=(5, 1.7), xytext=(1.5, 2.3),
               arrowprops=dict(arrowstyle='->', color='#333333', lw=1.5))
    ax.annotate('', xy=(5, 1.7), xytext=(5, 2.3),
               arrowprops=dict(arrowstyle='->', color='#333333', lw=1.5))
    ax.annotate('', xy=(5, 1.7), xytext=(8.5, 2.3),
               arrowprops=dict(arrowstyle='->', color='#333333', lw=1.5))

    plt.tight_layout()
    path = 'docs/images/multi_agent.png'
    os.makedirs('docs/images', exist_ok=True)
    plt.savefig(path, dpi=150, bbox_inches='tight', facecolor='white')
    plt.close()
    return path

def create_sse_streaming_diagram():
    """Create SSE streaming state machine diagram"""
    fig, ax = plt.subplots(1, 1, figsize=(10, 5))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 6)
    ax.axis('off')
    ax.set_title('SSE流式响应状态机设计', fontsize=14, fontweight='bold', pad=20)

    # Draw states
    states = [
        (2, 3, 'idle', '#F0F0F0'),
        (5, 3, 'streaming', '#FFD700'),
        (8, 3, 'complete', '#90EE90'),
    ]

    for x, y, text, color in states:
        circle = plt.Circle((x, y), 0.8, facecolor=color, edgecolor='#333333', linewidth=2)
        ax.add_patch(circle)
        ax.text(x, y, text, ha='center', va='center', fontsize=11, fontweight='bold')

    # Draw error state
    error = plt.Circle((8, 1), 0.6, facecolor='#FF6B6B', edgecolor='#333333', linewidth=2)
    ax.add_patch(error)
    ax.text(8, 1, 'error', ha='center', va='center', fontsize=10, fontweight='bold')

    # Draw arrows
    ax.annotate('', xy=(3.5, 3), xytext=(2.8, 3),
               arrowprops=dict(arrowstyle='->', color='#333333', lw=2))
    ax.annotate('', xy=(6.5, 3), xytext=(5.2, 3),
               arrowprops=dict(arrowstyle='->', color='#333333', lw=2))
    ax.annotate('', xy=(8, 1.8), xytext=(8, 3.8),
               arrowprops=dict(arrowstyle='->', color='#FF6B6B', lw=2))

    # Labels
    ax.text(4.25, 3.4, 'start', fontsize=9, style='italic')
    ax.text(5.75, 3.4, 'done', fontsize=9, style='italic')
    ax.text(8.3, 2.8, 'error', fontsize=9, style='italic', color='#FF6B6B')

    # Draw event descriptions
    events_text = '''
    双重发送机制：
    • row事件 → 边收边渲染
    • done事件 → 完整数据
    '''
    ax.text(0.5, 0.8, events_text, fontsize=9, va='top')

    plt.tight_layout()
    path = 'docs/images/sse_streaming.png'
    os.makedirs('docs/images', exist_ok=True)
    plt.savefig(path, dpi=150, bbox_inches='tight', facecolor='white')
    plt.close()
    return path

def create_reliability_diagram():
    """Create system reliability design diagram"""
    fig, ax = plt.subplots(1, 1, figsize=(10, 6))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 8)
    ax.axis('off')
    ax.set_title('系统可靠性设计 - 多级降级机制', fontsize=14, fontweight='bold', pad=20)

    # Draw the cascade
    levels = [
        (5, 7, 'LLM', '#87CEEB', '主力'),
        (5, 5.5, '规则引擎', '#FFD700', '一级降级'),
        (5, 4, '预设查询', '#FFA500', '二级降级'),
        (5, 2.5, '友好错误', '#FF6B6B', '最终兜底'),
    ]

    for x, y, text, color, note in levels:
        box = FancyBboxPatch((x-1.5, y-0.5), 3, 1,
                            boxstyle="round,pad=0.05,rounding_size=0.15",
                            facecolor=color, edgecolor='#333333', linewidth=2)
        ax.add_patch(box)
        ax.text(x, y+0.15, text, ha='center', va='center', fontsize=11, fontweight='bold')
        ax.text(x, y-0.25, note, ha='center', va='center', fontsize=8, style='italic')

    # Draw arrows
    for y in [6.25, 4.75, 3.25]:
        ax.annotate('', xy=(5, y-0.5), xytext=(5, y),
                   arrowprops=dict(arrowstyle='->', color='#333333', lw=2))

    # Draw formula
    formula = '''
    可用率 = 1 - (LLM失败率 × 降级失败率)
    示例: 1 - (15% × 5%) = 99.25%
    '''
    ax.text(0.5, 1.5, formula, fontsize=9, va='top', family='monospace')

    plt.tight_layout()
    path = 'docs/images/reliability.png'
    os.makedirs('docs/images', exist_ok=True)
    plt.savefig(path, dpi=150, bbox_inches='tight', facecolor='white')
    plt.close()
    return path

def create_skill_index_diagram():
    """Create Skill index table diagram"""
    fig, ax = plt.subplots(1, 1, figsize=(12, 8))
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 10)
    ax.axis('off')
    ax.set_title('LLM-Based SDD 可复用Skill索引', fontsize=14, fontweight='bold', pad=20)

    # Table data
    headers = ['阶段', 'Skill', '解决的问题']
    data = [
        ['技术选型', 'Skill 1: 技术选型三问法', '决策框架混乱'],
        ['技术选型', 'Skill 8: 多LLM Provider抽象', '供应商锁定'],
        ['意图识别', 'Skill 2: LLM与规则引擎级联', '不可解释性'],
        ['知识管理', 'Skill 3: RAG知识增强四步法', 'LLM记性不好'],
        ['开发调试', 'Skill 4: LLM异常排查三板斧', '输出异常定位'],
        ['开发调试', 'Skill 5: Prompt模板安全写法', '占位符冲突'],
        ['安全合规', 'Skill 6: LLM输出安全校验', 'SQL注入/全表扫描'],
        ['架构设计', 'Skill 7: Multi-Agent协作三问', 'Agent编排'],
        ['架构设计', 'Skill 9: 流式响应状态机', 'SSE完整性'],
        ['可靠性', 'Skill 10: 系统可靠性自检清单', '不知道测什么'],
        ['可靠性', 'Skill 11: LLM功能上线前检查', '上线遗漏'],
    ]

    # Draw header
    header_y = 8.5
    col_widths = [2, 4.5, 3]
    col_starts = [0.5, 2.5, 7]

    for i, (header, start, width) in enumerate(zip(headers, col_starts, col_widths)):
        rect = plt.Rectangle((start, header_y), width-0.1, 0.8,
                             facecolor='#4A90D9', edgecolor='#333333', linewidth=1)
        ax.add_patch(rect)
        ax.text(start + width/2, header_y+0.4, header, ha='center', va='center',
               fontsize=10, fontweight='bold', color='white')

    # Draw rows
    for row_idx, row in enumerate(data):
        y = header_y - (row_idx + 1) * 0.7
        color = '#F5F5F5' if row_idx % 2 == 0 else 'white'

        x_pos = col_starts[0]
        for col_idx, (cell, width) in enumerate(zip(row, col_widths)):
            rect = plt.Rectangle((x_pos, y), width-0.1, 0.65,
                                 facecolor=color, edgecolor='#CCCCCC', linewidth=0.5)
            ax.add_patch(rect)
            ax.text(x_pos + width/2, y + 0.325, cell, ha='center', va='center',
                   fontsize=8)

            if col_idx == 0:
                # Color code by category
                if '技术选型' in cell:
                    rect.set_facecolor('#FFE4B5')
                elif '意图识别' in cell:
                    rect.set_facecolor('#B0E0E6')
                elif '知识管理' in cell:
                    rect.set_facecolor('#98FB98')
                elif '开发调试' in cell:
                    rect.set_facecolor('#DDA0DD')
                elif '安全' in cell:
                    rect.set_facecolor('#FF6B6B')
                elif '架构' in cell:
                    rect.set_facecolor('#87CEEB')
                elif '可靠性' in cell:
                    rect.set_facecolor('#FFD700')

            x_pos += width

    plt.tight_layout()
    path = 'docs/images/skill_index.png'
    os.makedirs('docs/images', exist_ok=True)
    plt.savefig(path, dpi=150, bbox_inches='tight', facecolor='white')
    plt.close()
    return path

def create_llm_evolution_diagram():
    """Create LLM evolution roadmap diagram"""
    fig, ax = plt.subplots(1, 1, figsize=(12, 6))
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 8)
    ax.axis('off')
    ax.set_title('LLM软件工程演进路线图', fontsize=14, fontweight='bold', pad=20)

    # Draw levels
    levels = [
        (2, 6, 'Level 1\nLLM as Copilot', '#87CEEB', 'GitHub Copilot模式\n人类工程师+AI辅助'),
        (5, 4.5, 'Level 2\nLLM as Implementer', '#98FB98', 'OperatorBoard NL2SQL\n人类设计师+AI生成'),
        (8, 3, 'Level 3\nLLM as Agent', '#FFD700', 'Multi-Agent协作\n人类Orchestrator'),
        (11, 1.5, 'Level 4\nLLM as Architect', '#DDA0DD', 'AI自主完成全流程\n人类仅审核目标'),
    ]

    for x, y, text, color, desc in levels:
        box = FancyBboxPatch((x-1.5, y-1), 3, 2,
                            boxstyle="round,pad=0.05,rounding_size=0.15",
                            facecolor=color, edgecolor='#333333', linewidth=2)
        ax.add_patch(box)
        ax.text(x, y+0.3, text, ha='center', va='center', fontsize=10, fontweight='bold')
        ax.text(x, y-0.5, desc, ha='center', va='center', fontsize=7, style='italic')

    # Draw arrows
    for i in range(len(levels)-1):
        x1 = levels[i][0] + 1.5
        x2 = levels[i+1][0] - 1.5
        y1 = levels[i][1]
        y2 = levels[i+1][1]
        ax.annotate('', xy=(x2, y2), xytext=(x1, y1),
                   arrowprops=dict(arrowstyle='->', color='#333333', lw=2))

    # Current position marker
    ax.text(6.5, 5.5, '← 我们在这里', fontsize=10, style='italic', color='#FF6B6B', fontweight='bold')

    plt.tight_layout()
    path = 'docs/images/llm_evolution.png'
    os.makedirs('docs/images', exist_ok=True)
    plt.savefig(path, dpi=150, bbox_inches='tight', facecolor='white')
    plt.close()
    return path

def create_4p1_views_summary():
    """Create 4+1 architecture views summary diagram"""
    fig, axes = plt.subplots(2, 3, figsize=(15, 10))
    fig.suptitle('4+1 架构视图概览', fontsize=16, fontweight='bold', y=0.98)

    views = [
        ('场景视图', '功能需求\n用户交互', '#FFE4E1'),
        ('逻辑视图', '功能结构\n模块设计', '#E6E6FA'),
        ('进程视图', '并发异步\n执行流程', '#B0E0E6'),
        ('部署视图', '物理部署\n基础设施', '#98FB98'),
        ('开发视图', '代码组织\n依赖管理', '#FFD700'),
        ('技术架构', '技术选型\n系统设计', '#DDA0DD'),
    ]

    for idx, (title, content, color) in enumerate(views):
        ax = axes[idx // 3, idx % 3]
        ax.set_xlim(0, 4)
        ax.set_ylim(0, 3)
        ax.axis('off')

        box = FancyBboxPatch((0.2, 0.3), 3.6, 2.4,
                            boxstyle="round,pad=0.1,rounding_size=0.2",
                            facecolor=color, edgecolor='#333333', linewidth=2)
        ax.add_patch(box)
        ax.text(2, 2.3, title, ha='center', va='center', fontsize=14, fontweight='bold')
        ax.text(2, 1.3, content, ha='center', va='center', fontsize=11)

        # Add view number
        ax.text(0.4, 2.5, f'{idx+1}', fontsize=20, fontweight='bold', color='#666')

    # Hide the 6th subplot (we only have 5 views + technical)
    axes[1, 2].axis('off')

    plt.tight_layout()
    path = 'docs/images/4p1_views.png'
    os.makedirs('docs/images', exist_ok=True)
    plt.savefig(path, dpi=150, bbox_inches='tight', facecolor='white')
    plt.close()
    return path

def create_architecture_flow_diagram():
    """Create system architecture flow diagram"""
    fig, ax = plt.subplots(1, 1, figsize=(14, 10))
    ax.set_xlim(0, 14)
    ax.set_ylim(0, 12)
    ax.axis('off')
    ax.set_title('OperatorBoard 系统架构全景图', fontsize=14, fontweight='bold', pad=20)

    # Draw components
    components = [
        # Row 1: Frontend
        (7, 11, 'agent-app\n(React)', '#87CEEB', 2, 1),

        # Row 2: API Proxy
        (7, 9, 'API Proxy\n(Node.js)', '#98FB98', 2, 1),

        # Row 3: Agents
        (3, 7, 'operator-agent\n(Python)', '#FFD700', 2, 1.5),
        (7, 7, 'predict-agent', '#FFD700', 2, 1),
        (11, 7, 'auth-agent', '#FFD700', 2, 1),

        # Row 4: Services
        (3, 4.5, 'Java NL2SQL\nService', '#DDA0DD', 2.5, 1.5),
        (7, 4.5, 'SQLCoder\n(LLM)', '#FFA500', 2, 1),
        (11, 4.5, 'MySQL\nDatabase', '#FF6B6B', 2, 1.5),
    ]

    for x, y, text, color, w, h in components:
        box = FancyBboxPatch((x-w/2, y-h/2), w, h,
                            boxstyle="round,pad=0.05,rounding_size=0.15",
                            facecolor=color, edgecolor='#333333', linewidth=2)
        ax.add_patch(box)
        ax.text(x, y, text, ha='center', va='center', fontsize=9, fontweight='bold')

    # Draw arrows (vertical)
    arrow_props = dict(arrowstyle='->', color='#333333', lw=1.5)

    # Frontend to Proxy
    ax.annotate('', xy=(7, 9.5), xytext=(7, 10.5), arrowprops=arrow_props)

    # Proxy to Agents
    ax.annotate('', xy=(3, 7.8), xytext=(5.5, 8.5), arrowprops=arrow_props)
    ax.annotate('', xy=(7, 7.8), xytext=(7, 8.5), arrowprops=arrow_props)
    ax.annotate('', xy=(11, 7.8), xytext=(8.5, 8.5), arrowprops=arrow_props)

    # Agents to Services
    ax.annotate('', xy=(3, 5.3), xytext=(3, 6.2), arrowprops=arrow_props)
    ax.annotate('', xy=(7, 5.3), xytext=(7, 6.2), arrowprops=arrow_props)
    ax.annotate('', xy=(11, 5.3), xytext=(11, 6.2), arrowprops=arrow_props)

    # SQLCoder to NL2SQL Service
    ax.annotate('', xy=(5.5, 4.5), xytext=(6, 4.5), arrowprops=arrow_props)

    # NL2SQL to Database
    ax.annotate('', xy=(9.5, 4.5), xytext=(8.5, 4.5), arrowprops=arrow_props)

    # Add protocol labels
    ax.text(5, 9.2, 'HTTP', fontsize=7, style='italic')
    ax.text(5.2, 7.5, 'HTTP', fontsize=7, style='italic')
    ax.text(9, 7.5, 'MCP', fontsize=7, style='italic')
    ax.text(3.5, 5.8, 'REST', fontsize=7, style='italic')
    ax.text(8.2, 4.7, '/v1/completions', fontsize=6, style='italic')

    plt.tight_layout()
    path = 'docs/images/architecture_flow.png'
    os.makedirs('docs/images', exist_ok=True)
    plt.savefig(path, dpi=150, bbox_inches='tight', facecolor='white')
    plt.close()
    return path

def create_prompt_contract_diagram():
    """Create Prompt as Design Contract diagram"""
    fig, ax = plt.subplots(1, 1, figsize=(10, 6))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 8)
    ax.axis('off')
    ax.set_title('Prompt即设计契约', fontsize=14, fontweight='bold', pad=20)

    # Draw prompt as contract
    contract = FancyBboxPatch((2, 5), 6, 1.5,
                              boxstyle="round,pad=0.05,rounding_size=0.1",
                              facecolor='#FFE4B5', edgecolor='#FF8C00', linewidth=3)
    ax.add_patch(contract)
    ax.text(5, 6, 'Prompt = 设计契约', ha='center', va='center', fontsize=14, fontweight='bold')
    ax.text(5, 5.4, '版本号 | 变更记录 | 回归测试', ha='center', va='center', fontsize=9)

    # Draw comparisons
    code_contract = [
        (2, 3, '代码契约', '#E6E6FA', '违反 -> 编译报错'),
        (5, 3, 'Prompt契约', '#B0E0E6', '违反 -> 可能不知道'),
        (8, 3, 'API契约', '#98FB98', '违反 -> 接口报错'),
    ]

    for x, y, text, color, note in code_contract:
        box = FancyBboxPatch((x-1.2, y-0.6), 2.4, 1.4,
                            boxstyle="round,pad=0.05,rounding_size=0.1",
                            facecolor=color, edgecolor='#333333', linewidth=2)
        ax.add_patch(box)
        ax.text(x, y+0.3, text, ha='center', va='center', fontsize=10, fontweight='bold')
        ax.text(x, y-0.2, note, ha='center', va='center', fontsize=8, style='italic')

    # Key insight
    insight = FancyBboxPatch((1, 0.5), 8, 1,
                             boxstyle="round,pad=0.05,rounding_size=0.1",
                             facecolor='#FF6B6B', edgecolor='#333333', linewidth=2)
    ax.add_patch(insight)
    ax.text(5, 1, '关键洞察: Prompt契约需要像API契约一样管理', ha='center', va='center', fontsize=9, fontweight='bold', color='white')

    # Arrows from contract to comparisons
    ax.annotate('', xy=(2, 3.6), xytext=(4, 5),
               arrowprops=dict(arrowstyle='->', color='#666', lw=1))
    ax.annotate('', xy=(5, 3.6), xytext=(5, 5),
               arrowprops=dict(arrowstyle='->', color='#666', lw=1))
    ax.annotate('', xy=(8, 3.6), xytext=(6, 5),
               arrowprops=dict(arrowstyle='->', color='#666', lw=1))

    plt.tight_layout()
    path = 'docs/images/prompt_contract.png'
    os.makedirs('docs/images', exist_ok=True)
    plt.savefig(path, dpi=150, bbox_inches='tight', facecolor='white')
    plt.close()
    return path


def create_scenario_view_diagram():
    """Create Scenario View (Use Case diagram)"""
    fig, ax = plt.subplots(1, 1, figsize=(12, 8))
    ax.set_xlim(0, 14)
    ax.set_ylim(0, 10)
    ax.axis('off')
    ax.set_title('场景视图 - Scenario View', fontsize=14, fontweight='bold', pad=20)

    # Draw actors (stick figures simplified as circles)
    actors = [
        (1, 7.5, 'User', '#87CEEB'),
        (1, 5, 'Dashboard', '#98FB98'),
        (1, 2.5, 'NL2SQL\nService', '#DDA0DD'),
        (13, 5, 'MySQL', '#FF6B6B'),
        (13, 2.5, 'SQLCoder', '#FFD700'),
    ]

    for x, y, text, color in actors:
        circle = plt.Circle((x, y), 0.6, facecolor=color, edgecolor='#333333', linewidth=2)
        ax.add_patch(circle)
        ax.text(x, y, text, ha='center', va='center', fontsize=8, fontweight='bold')

    # Draw use cases
    use_cases = [
        (7, 8, 'UC-001 NL2SQL查询', '#FFE4E1'),
        (7, 6.5, 'UC-002 运营商查询', '#FFE4E1'),
        (7, 5, 'UC-003 指标查询', '#FFE4E1'),
        (7, 3.5, 'UC-004 历史查询', '#FFE4E1'),
        (10, 2, 'UC-005 数据看板', '#FFE4E1'),
    ]

    for x, y, text, color in use_cases:
        box = FancyBboxPatch((x-1.5, y-0.5), 3, 1,
                            boxstyle="round,pad=0.05,rounding_size=0.1",
                            facecolor=color, edgecolor='#333333', linewidth=2)
        ax.add_patch(box)
        ax.text(x, y, text, ha='center', va='center', fontsize=9, fontweight='bold')

    # Draw relationships (simplified lines)
    # User to use cases
    ax.plot([1.6, 5.5], [7.5, 8], 'k-', lw=1)
    ax.plot([1.6, 5.5], [7.5, 6.5], 'k-', lw=1)
    ax.plot([1.6, 5.5], [7.5, 5], 'k-', lw=1)
    ax.plot([1.6, 5.5], [7.5, 3.5], 'k-', lw=1)

    # Use cases to services
    ax.plot([8.5, 12], [5, 13], 'k--', lw=1)  # to MySQL
    ax.plot([8.5, 10], [3.5, 12.5], 'k--', lw=1)  # to SQLCoder

    plt.tight_layout()
    path = 'docs/images/scenario_view.png'
    os.makedirs('docs/images', exist_ok=True)
    plt.savefig(path, dpi=150, bbox_inches='tight', facecolor='white')
    plt.close()
    return path


def create_process_view_diagram():
    """Create Process View (Activity diagram)"""
    fig, ax = plt.subplots(1, 1, figsize=(12, 10))
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 12)
    ax.axis('off')
    ax.set_title('进程视图 - Process View', fontsize=14, fontweight='bold', pad=20)

    # Draw activity boxes
    activities = [
        (3, 10, 'User Request\nInput', '#FFE4E1'),
        (3, 8, 'Chat/Dashboard\nDecision', '#E6E6FA'),
        (6, 8, 'Intent\nDetection', '#B0E0E6'),
        (9, 8, 'Route to\nService', '#98FB98'),
        (3, 5, 'Parallel\nQueries', '#FFD700'),
        (6, 5, 'NL2SQL\nProcessing', '#DDA0DD'),
        (9, 5, 'SQL\nExecution', '#FFA500'),
        (6, 2, 'Response\nSerialization', '#87CEEB'),
        (3, 2, 'Visualization\n(Recharts)', '#C5E0B4'),
    ]

    for x, y, text, color in activities:
        box = FancyBboxPatch((x-1.5, y-0.7), 3, 1.6,
                            boxstyle="round,pad=0.05,rounding_size=0.1",
                            facecolor=color, edgecolor='#333333', linewidth=2)
        ax.add_patch(box)
        ax.text(x, y, text, ha='center', va='center', fontsize=9, fontweight='bold')

    # Draw arrows
    arrows = [
        (3, 9.3, 3, 8.7),
        (4.5, 8, 4.5, 8),
        (6, 8.7, 6, 8.3),
        (7.5, 8, 7.5, 8),
        (9, 8.7, 9, 8.3),
        (9, 7.3, 9, 5.7),
        (8, 5, 7, 5),
        (6, 5.7, 6, 5.3),
        (6, 4.3, 6, 2.7),
        (5, 2, 4, 2),
        (2, 2.7, 2, 2.3),
    ]

    for i in range(len(arrows)-1):
        x1, y1 = arrows[i][0], arrows[i][1]
        x2, y2 = arrows[i+1][0], arrows[i+1][1]
        ax.annotate('', xy=(x2, y2), xytext=(x1, y1),
                   arrowprops=dict(arrowstyle='->', color='#333333', lw=1.5))

    # Add fork/join notation
    ax.text(3, 6.5, 'fork', fontsize=8, style='italic')
    ax.text(3, 4.5, 'join', fontsize=8, style='italic')

    plt.tight_layout()
    path = 'docs/images/process_view.png'
    os.makedirs('docs/images', exist_ok=True)
    plt.savefig(path, dpi=150, bbox_inches='tight', facecolor='white')
    plt.close()
    return path


def create_deployment_view_diagram():
    """Create Deployment View diagram"""
    fig, ax = plt.subplots(1, 1, figsize=(12, 8))
    ax.set_xlim(0, 14)
    ax.set_ylim(0, 10)
    ax.axis('off')
    ax.set_title('部署视图 - Deployment View', fontsize=14, fontweight='bold', pad=20)

    # Draw nodes
    nodes = [
        (2, 8, 'End Users\n(REST)', '#87CEEB'),
        (2, 5.5, 'API Gateway\nLoad Balancer', '#98FB98'),
        (5, 5.5, 'Agent N\n(Python)', '#FFD700'),
        (8, 5.5, 'Java Services\n(MVC+CQRS)', '#DDA0DD'),
        (11, 5.5, 'MySQL\nDatabase', '#FF6B6B'),
        (11, 2.5, 'SQLCoder\nLLM', '#FFA500'),
    ]

    for x, y, text, color in nodes:
        box = FancyBboxPatch((x-1.5, y-0.8), 3, 1.8,
                            boxstyle="round,pad=0.05,rounding_size=0.15",
                            facecolor=color, edgecolor='#333333', linewidth=2)
        ax.add_patch(box)
        ax.text(x, y, text, ha='center', va='center', fontsize=9, fontweight='bold')

    # Draw connections
    connections = [
        (2, 7.2, 2, 6.3),
        (3.5, 5.5, 3.5, 5.5),
        (5, 6.3, 5, 6.3),
        (6.5, 5.5, 6.5, 5.5),
        (8, 6.3, 8, 6.3),
        (9.5, 5.5, 9.5, 5.5),
        (11, 6.3, 11, 4.3),
        (9.5, 2.5, 10.5, 2.5),
    ]

    for i in range(len(connections)-1):
        x1, y1 = connections[i][0], connections[i][1]
        x2, y2 = connections[i+1][0], connections[i+1][1]
        ax.annotate('', xy=(x2, y2), xytext=(x1, y1),
                   arrowprops=dict(arrowstyle='->', color='#333333', lw=1.5))

    # Labels
    ax.text(2.5, 7.8, 'HTTPS', fontsize=7, style='italic')
    ax.text(4.2, 5.8, 'HTTP', fontsize=7, style='italic')
    ax.text(7.2, 5.8, 'HTTP', fontsize=7, style='italic')
    ax.text(10.2, 5.8, 'JDBC', fontsize=7, style='italic')
    ax.text(10.2, 3.8, 'REST', fontsize=7, style='italic')

    plt.tight_layout()
    path = 'docs/images/deployment_view.png'
    os.makedirs('docs/images', exist_ok=True)
    plt.savefig(path, dpi=150, bbox_inches='tight', facecolor='white')
    plt.close()
    return path


def create_development_view_diagram():
    """Create Development View diagram"""
    fig, ax = plt.subplots(1, 1, figsize=(12, 8))
    ax.set_xlim(0, 14)
    ax.set_ylim(0, 10)
    ax.axis('off')
    ax.set_title('开发视图 - Development View', fontsize=14, fontweight='bold', pad=20)

    # Draw modules
    modules = [
        # Row 1
        (2, 8, 'agent-app', '#87CEEB', 'React\nVite'),
        (5, 8, 'agent-framework', '#98FB98', 'Core\nPython'),
        (8, 8, 'operator-agent', '#FFD700', 'Business\nPython'),
        (11, 8, 'operator-service', '#DDA0DD', 'Java\nSpring'),
        # Row 2
        (2, 5, 'operator-dashboard', '#FFA500', 'Components'),
        (5, 5, 'BaseAgent\nBaseTool', '#C5E0B4', 'Framework'),
        (8, 5, 'Tools\nSkills\nMCP', '#E6E6FA', 'Capabilities'),
        (11, 5, 'Controller\nService\nRepository', '#FFE4E1', 'MVC+CQRS'),
        # Row 3
        (3.5, 2, 'npm packages', '#87CEEB', 'deps'),
        (7, 2, 'pip packages', '#98FB98', 'deps'),
        (10.5, 2, 'Maven\nJAR', '#DDA0DD', 'deps'),
    ]

    for x, y, text, color, note in modules:
        box = FancyBboxPatch((x-1.5, y-0.7), 3, 1.6,
                            boxstyle="round,pad=0.05,rounding_size=0.1",
                            facecolor=color, edgecolor='#333333', linewidth=2)
        ax.add_patch(box)
        ax.text(x, y+0.3, text, ha='center', va='center', fontsize=9, fontweight='bold')
        ax.text(x, y-0.4, note, ha='center', va='center', fontsize=7, style='italic')

    # Draw module dependencies
    dep_arrows = [
        ((2, 7.3), (5, 7.3)),
        ((5, 7.3), (8, 7.3)),
        ((8, 7.3), (11, 7.3)),
        ((2, 4.3), (2, 2.7)),
        ((5, 4.3), (5, 2.7)),
        ((11, 4.3), (11, 2.7)),
    ]

    for start, end in dep_arrows:
        ax.annotate('', xy=end, xytext=start,
                   arrowprops=dict(arrowstyle='->', color='#666', lw=1, linestyle='dashed'))

    plt.tight_layout()
    path = 'docs/images/development_view.png'
    os.makedirs('docs/images', exist_ok=True)
    plt.savefig(path, dpi=150, bbox_inches='tight', facecolor='white')
    plt.close()
    return path

def create_all_diagrams():
    """Generate all diagrams and return their paths"""
    diagrams = {}

    print("Generating diagrams...")

    print("  - SDD Evolution diagram...")
    diagrams['sdd_evolution'] = create_sdd_evolution_diagram()

    print("  - LLM Provider diagram...")
    diagrams['llm_provider'] = create_llm_provider_diagram()

    print("  - Intent Detection diagram...")
    diagrams['intent_detection'] = create_intent_detection_diagram()

    print("  - RAG Knowledge diagram...")
    diagrams['rag_knowledge'] = create_rag_knowledge_diagram()

    print("  - SQL Safety diagram...")
    diagrams['sql_safety'] = create_sql_safety_diagram()

    print("  - Multi-Agent diagram...")
    diagrams['multi_agent'] = create_multi_agent_diagram()

    print("  - SSE Streaming diagram...")
    diagrams['sse_streaming'] = create_sse_streaming_diagram()

    print("  - Reliability diagram...")
    diagrams['reliability'] = create_reliability_diagram()

    print("  - Skill Index diagram...")
    diagrams['skill_index'] = create_skill_index_diagram()

    print("  - LLM Evolution diagram...")
    diagrams['llm_evolution'] = create_llm_evolution_diagram()

    print("  - 4+1 Views diagram...")
    diagrams['4p1_views'] = create_4p1_views_summary()

    print("  - Scenario View diagram...")
    diagrams['scenario_view'] = create_scenario_view_diagram()

    print("  - Process View diagram...")
    diagrams['process_view'] = create_process_view_diagram()

    print("  - Deployment View diagram...")
    diagrams['deployment_view'] = create_deployment_view_diagram()

    print("  - Development View diagram...")
    diagrams['development_view'] = create_development_view_diagram()

    print("  - Architecture Flow diagram...")
    diagrams['architecture_flow'] = create_architecture_flow_diagram()

    print("  - Prompt Contract diagram...")
    diagrams['prompt_contract'] = create_prompt_contract_diagram()

    print("All diagrams generated successfully!")
    return diagrams

def read_blog_content():
    """Read the blog content from markdown file"""
    blog_path = 'docs/技术博客-大模型SDD开发趟坑笔记.md'
    with open(blog_path, 'r', encoding='utf-8') as f:
        return f.read()

def create_word_document(diagrams):
    """Create Word document from blog content with diagrams"""
    doc = Document()

    # Set document properties
    doc.core_properties.title = '大模型时代软件开发者的第三次思考'
    doc.core_properties.author = 'OperatorBoard'

    # Add title
    title = doc.add_heading('大模型时代，软件开发者的第三次思考', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add subtitle
    subtitle = doc.add_paragraph('一位十余年软件工程实践者的 LLM-Based SDD 探索手记')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # spacing

    # Read and process blog content
    content = read_blog_content()

    # Process content section by section
    lines = content.split('\n')
    in_code_block = False
    code_content = []
    code_language = ""

    for line in lines:
        # Skip the title and subtitle (already added)
        if line.startswith('# 大模型时代') or line.startswith('> 一位十余年'):
            continue

        # Handle headings
        if line.startswith('## '):
            doc.add_heading(line[4:], 2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], 3)
        elif line.startswith('#### '):
            doc.add_heading(line[4:], 4)

        # Handle horizontal rule
        elif line.startswith('---'):
            doc.add_paragraph('─' * 50)

        # Handle code blocks
        elif line.startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_language = line[3:].strip()
                code_content = []
            else:
                in_code_block = False
                # Add code block
                code_text = '\n'.join(code_content)
                add_code_block(doc, code_text, code_language)
        elif in_code_block:
            code_content.append(line)

        # Handle tables (simple approach)
        elif line.startswith('|'):
            # Skip table formatting lines
            if set(line.replace('-', '').replace('|', '').replace(':', '').replace(' ', '')) == set():
                continue
            # Parse and add table
            parts = [p.strip() for p in line.split('|')[1:-1]]
            if parts and any(parts):
                para = doc.add_paragraph()
                para.add_run(' | '.join(parts)).font.size = Pt(9)

        # Handle blockquotes
        elif line.startswith('>'):
            para = doc.add_paragraph()
            run = para.add_run(line[1:].strip())
            run.italic = True

        # Handle bullet points
        elif line.startswith('- '):
            doc.add_paragraph(line[2:], style='List Bullet')

        # Handle numbered lists
        elif line.strip() and line[0].isdigit() and '. ' in line[:5]:
            doc.add_paragraph(line, style='List Number')

        # Handle bold text
        elif '**' in line:
            parts = line.split('**')
            para = doc.add_paragraph()
            for i, part in enumerate(parts):
                run = para.add_run(part)
                if i % 2 == 1:
                    run.bold = True
        else:
            # Regular paragraph
            if line.strip():
                doc.add_paragraph(line)

    return doc

def add_diagrams_to_document(doc, diagrams):
    """Add diagrams at appropriate positions in the document"""

    # This would require re-processing the content and inserting images
    # For simplicity, we'll add diagrams at the end

    doc.add_page_break()
    doc.add_heading('插图索引', 1)

    # Add each diagram
    diagram_descriptions = [
        ('sdd_evolution', 'SDD三次认知迭代', '展示了从2019年DDD到2024年LLM-Based SDD的演进过程'),
        ('llm_provider', '多LLM Provider抽象设计', '展示了如何设计可切换的LLM Provider架构'),
        ('intent_detection', '意图识别级联设计', '展示了LLM与规则引擎的协作模式'),
        ('rag_knowledge', 'RAG知识增强四步法', '展示了知识分层和检索增强流程'),
        ('sql_safety', 'LLM输出安全校验框架', '展示了SQL安全验证的三层架构'),
        ('multi_agent', 'Multi-Agent协作模式', '展示了多Agent协作的问题解决流程'),
        ('sse_streaming', 'SSE流式响应状态机', '展示了流式响应的状态管理'),
        ('reliability', '系统可靠性设计', '展示了多级降级机制'),
        ('skill_index', 'Skill索引表', '汇总了11个可复用Skill'),
        ('llm_evolution', 'LLM软件工程演进路线图', '展示了从Copilot到Architect的演进'),
        ('4p1_views', '4+1架构视图概览', '展示了5个架构视图'),
        ('architecture_flow', '系统架构全景图', '展示了完整的系统架构'),
        ('prompt_contract', 'Prompt即设计契约', '展示了Prompt作为设计契约的概念'),
        ('scenario_view', '场景视图 - Scenario View', '展示了用例图和参与者关系'),
        ('process_view', '进程视图 - Process View', '展示了活动流程和异步处理'),
        ('deployment_view', '部署视图 - Deployment View', '展示了部署拓扑和节点关系'),
        ('development_view', '开发视图 - Development View', '展示了代码组织和模块依赖'),
    ]

    for key, title, desc in diagram_descriptions:
        if key in diagrams:
            doc.add_heading(title, 2)
            doc.add_paragraph(desc)
            try:
                add_image(doc, diagrams[key], width=Inches(5.5))
            except Exception as e:
                print(f"Warning: Could not add image {key}: {e}")

def main():
    """Main function to generate the Word document"""
    print("=" * 60)
    print("Generating Word Document from Technical Blog")
    print("=" * 60)

    # Create diagrams
    print("\n[1/3] Generating diagrams...")
    diagrams = create_all_diagrams()

    # Create document
    print("\n[2/3] Creating Word document...")
    doc = create_word_document(diagrams)

    # Add diagrams section
    print("\n[3/3] Adding diagrams to document...")
    add_diagrams_to_document(doc, diagrams)

    # Save document
    output_path = 'docs/技术博客-大模型SDD开发趟坑笔记.docx'
    doc.save(output_path)
    print(f"\nDocument saved to: {output_path}")

    print("\n" + "=" * 60)
    print("Generation complete!")
    print("=" * 60)

if __name__ == '__main__':
    main()
