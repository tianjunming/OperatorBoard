"""Generate a single merged Word document from all markdown documentation."""

import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from pathlib import Path

DOCS_DIR = Path(__file__).parent / "views"
OUTPUT_FILE = DOCS_DIR / "output" / "OperatorBoard架构设计文档.docx"


def create_word_doc(md_file, doc, add_heading=True):
    """Add content from a markdown file to the document."""
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')
    i = 0
    in_table = False
    table_lines = []

    while i < len(lines):
        line = lines[i]

        # Handle code blocks
        if line.strip().startswith('```'):
            if in_table:
                add_table_to_doc(doc, table_lines)
                table_lines = []
                in_table = False

            # Find the full code block
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1

            # Add code block
            if code_lines:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.3)
                run = p.add_run('\n'.join(code_lines))
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0, 100, 0)
            i += 1
            continue

        # Handle tables
        if line.strip().startswith('|'):
            in_table = True
            table_lines.append(line)
        else:
            if in_table:
                add_table_to_doc(doc, table_lines)
                table_lines = []
                in_table = False

            # Handle headers
            if line.startswith('#### '):
                doc.add_heading(line[5:], level=4)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.startswith('## '):
                if add_heading:
                    doc.add_heading(line[3:], level=2)
            elif line.startswith('# '):
                if add_heading:
                    doc.add_heading(line[2:], level=1)
            elif line.strip():
                clean_line = line
                clean_line = re.sub(r'\*\*(.*?)\*\*', r'\1', clean_line)
                clean_line = re.sub(r'\*(.*?)\*', r'\1', clean_line)
                clean_line = re.sub(r'`(.*?)`', r'\1', clean_line)
                clean_line = re.sub(r'^-\s+', '• ', clean_line)

                p = doc.add_paragraph(clean_line)

        i += 1

    if in_table:
        add_table_to_doc(doc, table_lines)


def add_table_to_doc(doc, table_lines):
    """Add a table to the document from markdown table lines."""
    if not table_lines or len(table_lines) < 2:
        return

    headers = [cell.strip() for cell in table_lines[0].split('|')[1:-1]]

    if len(table_lines) > 1 and re.match(r'^\|[\s\-:|]+\|$', table_lines[1]):
        data_lines = table_lines[2:]
    else:
        data_lines = table_lines[1:]

    rows = []
    for line in data_lines:
        if line.strip():
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            rows.append(cells)

    if headers and rows:
        table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
        table.style = 'Table Grid'

        for j, header in enumerate(headers):
            cell = table.rows[0].cells[j]
            cell.text = header
            if cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].font.bold = True

        for row_idx, row_data in enumerate(rows):
            for col_idx, cell_text in enumerate(row_data):
                if col_idx < len(headers):
                    table.rows[row_idx + 1].cells[col_idx].text = cell_text

        doc.add_paragraph()


def main():
    """Generate the merged Word document."""
    OUTPUT_FILE.parent.mkdir(exist_ok=True)

    # Create new document
    doc = Document()

    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # Set default style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    # Document order
    doc_files = [
        ('00-index.md', '00-目录', True),
        ('01-scenario-view.md', '01-场景视图', False),
        ('02-logical-view.md', '02-逻辑视图', False),
        ('03-process-view.md', '03-进程视图', False),
        ('04-deployment-view.md', '04-部署视图', False),
        ('05-development-view.md', '05-开发视图', False),
        ('06-technical-architecture.md', '06-技术架构与选型', False),
    ]

    for md_file, title, is_first in doc_files:
        md_path = DOCS_DIR / md_file
        if md_path.exists():
            # Add page break for all except first
            if not is_first:
                doc.add_page_break()

            # Add title
            clean_title = title.replace('0', '').replace('-', ' ')
            doc.add_heading(clean_title, level=1)

            # Add content
            create_word_doc(md_path, doc, add_heading=False)
            print(f"Added: {md_file}")
        else:
            print(f"Warning: {md_path} not found")

    # Save document
    doc.save(OUTPUT_FILE)
    print(f"\nCreated merged document: {OUTPUT_FILE}")


if __name__ == '__main__':
    main()
